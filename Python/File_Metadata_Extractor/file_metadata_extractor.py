#!/usr/bin/env python3
"""
File Metadata Extractor - Comprehensive metadata extraction tool.

This module provides a comprehensive solution for extracting metadata from
various file types commonly used in digital asset management, forensic analysis,
and file organization workflows.

Supported File Types:
    Images:
        - JPEG (with full EXIF data extraction)
        - PNG (dimensions, transparency info)
        - TIFF (multi-page support, EXIF data)
        - BMP, GIF (basic properties)

    Audio Files:
        - MP3 (ID3 tags, bitrate, duration)
        - FLAC (lossless metadata)
        - WAV (audio properties)
        - M4A, AAC, OGG, WMA (various metadata)

    Video Files:
        - MP4, AVI, MKV (duration, bitrate)
        - MOV, WMV, FLV, WebM (basic properties)

    Documents:
        - PDF (page count, author, encryption status)
        - DOCX (word count, author, document properties)

Key Features:
    - Single file and batch processing
    - Multiple output formats (JSON, CSV)
    - Comprehensive error handling
    - Cross-platform compatibility
    - Security-conscious file path validation

Usage:
    python file_metadata_extractor.py /path/to/file.jpg
    python file_metadata_extractor.py /path/to/directory/ --recursive -o results.json

Dependencies:
    - Pillow: Image processing and EXIF extraction
    - mutagen: Audio and video metadata
    - PyPDF2: PDF document analysis
    - python-docx: Microsoft Word documents

Author: Created for Rotten-Scripts Repository
Date: October 2025
License: MIT (following repository standards)
"""

import os
import sys
import json
import csv
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List
import argparse

# Import libraries for metadata extraction with graceful fallback
# Each library is optional - the script will skip unsupported formats if missing

try:
    # PIL/Pillow for image processing and EXIF data extraction
    from PIL import Image
    from PIL.ExifTags import TAGS

    PIL_AVAILABLE = True
except ImportError:
    # Image processing will be disabled if PIL is not available
    PIL_AVAILABLE = False

try:
    # Mutagen for audio and video metadata extraction
    import mutagen
    from mutagen import File as MutagenFile

    MUTAGEN_AVAILABLE = True
except ImportError:
    # Audio/video metadata extraction will be disabled
    MUTAGEN_AVAILABLE = False

try:
    # PyPDF2 for PDF document metadata extraction
    import PyPDF2

    PYPDF2_AVAILABLE = True
except ImportError:
    # PDF metadata extraction will be disabled
    PYPDF2_AVAILABLE = False

try:
    # python-docx for Microsoft Word document processing
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    # DOCX metadata extraction will be disabled
    DOCX_AVAILABLE = False


class FileMetadataExtractor:
    """
    A comprehensive metadata extractor for various file types.

    This class provides methods to extract metadata from images, audio files,
    videos, and documents. It supports multiple file formats and provides
    both individual file analysis and batch processing capabilities.

    Supported formats:
        - Images: JPEG (EXIF), PNG, TIFF, BMP, GIF
        - Audio: MP3 (ID3), FLAC, WAV, M4A, AAC, OGG, WMA
        - Video: MP4, AVI, MKV, MOV, WMV, FLV, WebM
        - Documents: PDF, DOCX

    Attributes:
        supported_image_formats (set): Set of supported image file extensions
        supported_audio_formats (set): Set of supported audio file extensions
        supported_video_formats (set): Set of supported video file extensions
        supported_document_formats (set): Set of supported document file extensions
    """

    def __init__(self):
        """
        Initialize the metadata extractor with supported file formats.

        Sets up the file extension mappings for different media types.
        No external dependencies are required for initialization.
        """
        self.supported_image_formats = {
            ".jpg",
            ".jpeg",
            ".png",
            ".tiff",
            ".tif",
            ".bmp",
            ".gif",
        }
        self.supported_audio_formats = {
            ".mp3",
            ".flac",
            ".wav",
            ".m4a",
            ".aac",
            ".ogg",
            ".wma",
        }
        self.supported_video_formats = {
            ".mp4",
            ".avi",
            ".mkv",
            ".mov",
            ".wmv",
            ".flv",
            ".webm",
        }
        self.supported_document_formats = {".pdf", ".docx", ".doc"}

    @staticmethod
    def get_basic_file_info(file_path: str) -> Dict[str, Any]:
        """
        Extract basic file system information.

        Args:
            file_path (str): Path to the file to analyze

        Returns:
            Dict[str, Any]: Dictionary containing file system metadata including
                filename, size, timestamps, and extension
        """
        try:
            stat_info = os.stat(file_path)
            file_size = stat_info.st_size

            return {
                "filename": os.path.basename(file_path),
                "file_path": os.path.abspath(file_path),
                "file_size_bytes": file_size,
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "extension": Path(file_path).suffix.lower(),
                "created_time": datetime.fromtimestamp(stat_info.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
                "accessed_time": datetime.fromtimestamp(stat_info.st_atime).isoformat(),
            }
        except Exception as e:
            return {"error": f"Error getting basic file info: {str(e)}"}

    @staticmethod
    def extract_image_metadata(file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from image files including EXIF data.

        Args:
            file_path (str): Path to the image file

        Returns:
            Dict[str, Any]: Dictionary containing image metadata such as
                dimensions, format, EXIF data, and color information
        """
        if not PIL_AVAILABLE:
            return {
                "error": "PIL/Pillow not available. Install with: pip install Pillow"
            }

        try:
            with Image.open(file_path) as image:
                # Basic image info
                metadata = {
                    "width": image.width,
                    "height": image.height,
                    "format": image.format,
                    "mode": image.mode,
                    "has_transparency": image.mode in ("RGBA", "LA")
                    or "transparency" in image.info,
                }

                # EXIF data for JPEG images
                exif = getattr(image, "_getexif", lambda: None)()
                if exif is not None:
                    exif_data = {}

                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        try:
                            # Convert bytes to string if needed
                            if isinstance(value, bytes):
                                value = value.decode("utf-8", errors="ignore")
                            exif_data[tag] = str(value)
                        except (TypeError, ValueError):
                            exif_data[tag] = str(value)

                    metadata["exif"] = exif_data

                return metadata

        except Exception as e:
            return {"error": f"Error extracting image metadata: {str(e)}"}

    @staticmethod
    def extract_audio_metadata(file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from audio files including ID3 tags and technical info.

        Args:
            file_path (str): Path to the audio file to analyze

        Returns:
            Dict[str, Any]: Dictionary containing audio metadata including:
                - duration_seconds: Length of audio in seconds
                - bitrate: Audio bitrate in kbps
                - sample_rate: Sample rate in Hz
                - channels: Number of audio channels
                - title, artist, album: ID3 tag information
                - year, genre: Additional metadata fields

        Note:
            Requires mutagen library. Returns error message if not available.
        """
        if not MUTAGEN_AVAILABLE:
            return {"error": "Mutagen not available. Install with: pip install mutagen"}

        try:
            audio_file = MutagenFile(file_path)
            if audio_file is None:
                return {"error": "Unsupported audio format"}

            metadata = {}

            # Basic audio info
            if hasattr(audio_file, "info"):
                info = audio_file.info
                metadata.update(
                    {
                        "duration_seconds": getattr(info, "length", 0),
                        "bitrate": getattr(info, "bitrate", 0),
                        "sample_rate": getattr(info, "sample_rate", 0),
                        "channels": getattr(info, "channels", 0),
                    }
                )

                # Convert duration to readable format
                if metadata.get("duration_seconds"):
                    duration = int(metadata["duration_seconds"])
                    minutes, seconds = divmod(duration, 60)
                    metadata["duration_formatted"] = f"{minutes}:{seconds:02d}"

            # Tag information
            if audio_file.tags:
                tag_mapping = {
                    "TIT2": "title",
                    "TPE1": "artist",
                    "TALB": "album",
                    "TDRC": "year",
                    "TCON": "genre",
                    "TPE2": "album_artist",
                    "TRCK": "track_number",
                }

                for tag_key, readable_key in tag_mapping.items():
                    if tag_key in audio_file.tags:
                        metadata[readable_key] = str(audio_file.tags[tag_key][0])

                # Handle other common tag formats
                for key, value in audio_file.tags.items():
                    if isinstance(key, str) and key.lower() in [
                        "title",
                        "artist",
                        "album",
                        "date",
                    ]:
                        metadata[key.lower()] = (
                            str(value[0]) if isinstance(value, list) else str(value)
                        )

            return metadata

        except Exception as e:
            return {"error": f"Error extracting audio metadata: {str(e)}"}

    @staticmethod
    def extract_video_metadata(file_path: str) -> Dict[str, Any]:
        """
        Extract basic metadata from video files.

        Args:
            file_path (str): Path to the video file to analyze

        Returns:
            Dict[str, Any]: Dictionary containing video metadata including:
                - duration_seconds: Length of video in seconds
                - duration_formatted: Human-readable duration (HH:MM:SS)
                - bitrate: Video bitrate if available

        Note:
            Uses mutagen library for basic video metadata extraction.
            More detailed video analysis would require additional libraries.
        """
        if not MUTAGEN_AVAILABLE:
            return {"error": "Mutagen not available. Install with: pip install mutagen"}

        try:
            video_file = MutagenFile(file_path)
            if video_file is None:
                return {"error": "Unsupported video format or file"}

            metadata = {}

            # Basic video info
            if hasattr(video_file, "info"):
                info = video_file.info
                metadata.update(
                    {
                        "duration_seconds": getattr(info, "length", 0),
                        "bitrate": getattr(info, "bitrate", 0),
                    }
                )

                # Convert duration to readable format
                if metadata.get("duration_seconds"):
                    duration = int(metadata["duration_seconds"])
                    hours, remainder = divmod(duration, 3600)
                    minutes, seconds = divmod(remainder, 60)
                    if hours > 0:
                        metadata["duration_formatted"] = (
                            f"{hours}:{minutes:02d}:{seconds:02d}"
                        )
                    else:
                        metadata["duration_formatted"] = f"{minutes}:{seconds:02d}"

            return metadata

        except Exception as e:
            return {"error": f"Error extracting video metadata: {str(e)}"}

    @staticmethod
    def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from PDF documents.

        Args:
            file_path (str): Path to the PDF file to analyze

        Returns:
            Dict[str, Any]: Dictionary containing PDF metadata including:
                - page_count: Number of pages in the document
                - is_encrypted: Whether the PDF is password protected
                - title: Document title if available
                - author: Document author if available
                - subject: Document subject if available
                - creator: Application that created the PDF
                - producer: Application that produced the PDF
                - creation_date: When the document was created
                - modification_date: When the document was last modified

        Note:
            Requires PyPDF2 library for PDF processing.
        """
        if not PYPDF2_AVAILABLE:
            return {"error": "PyPDF2 not available. Install with: pip install PyPDF2"}

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                metadata = {
                    "page_count": len(pdf_reader.pages),
                    "is_encrypted": pdf_reader.is_encrypted,
                }

                # Extract document info
                if pdf_reader.metadata:
                    doc_info = pdf_reader.metadata
                    info_mapping = {
                        "/Title": "title",
                        "/Author": "author",
                        "/Subject": "subject",
                        "/Creator": "creator",
                        "/Producer": "producer",
                        "/CreationDate": "creation_date",
                        "/ModDate": "modification_date",
                    }

                    for pdf_key, readable_key in info_mapping.items():
                        if pdf_key in doc_info:
                            metadata[readable_key] = str(doc_info[pdf_key])

                return metadata

        except Exception as e:
            return {"error": f"Error extracting PDF metadata: {str(e)}"}

    @staticmethod
    def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from Microsoft Word DOCX documents.

        Args:
            file_path (str): Path to the DOCX file to analyze

        Returns:
            Dict[str, Any]: Dictionary containing DOCX metadata including:
                - title: Document title
                - author: Document author
                - subject: Document subject
                - keywords: Document keywords
                - comments: Document comments
                - created: Creation timestamp
                - modified: Last modification timestamp
                - last_modified_by: User who last modified the document
                - revision: Document revision number
                - paragraph_count: Number of paragraphs
                - table_count: Number of tables
                - word_count: Approximate word count

        Note:
            Requires python-docx library for DOCX processing.
        """
        if not DOCX_AVAILABLE:
            return {
                "error": "python-docx not available. Install with: pip install python-docx"
            }

        try:
            doc = Document(file_path)

            metadata = {}

            # Core properties
            if doc.core_properties:
                props = doc.core_properties
                metadata.update(
                    {
                        "title": props.title or "",
                        "author": props.author or "",
                        "subject": props.subject or "",
                        "keywords": props.keywords or "",
                        "comments": props.comments or "",
                        "created": props.created.isoformat() if props.created else "",
                        "modified": props.modified.isoformat()
                        if props.modified
                        else "",
                        "last_modified_by": props.last_modified_by or "",
                        "revision": props.revision or "",
                    }
                )

            # Document statistics
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)

            # Word count (approximate)
            word_count = 0
            for paragraph in doc.paragraphs:
                word_count += len(paragraph.text.split())
            metadata["word_count"] = word_count

            return metadata

        except Exception as e:
            return {"error": f"Error extracting DOCX metadata: {str(e)}"}

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract comprehensive metadata from a file based on its type.

        This is the main method that determines the file type and calls
        the appropriate extraction method. It combines basic file system
        information with format-specific metadata.

        Args:
            file_path (str): Path to the file to analyze

        Returns:
            Dict[str, Any]: Dictionary containing comprehensive metadata including:
                - Basic file info (size, dates, extension)
                - Format-specific metadata (EXIF, ID3 tags, etc.)
                - File type classification
                - Error information if extraction fails
        """
        # Validate and sanitize file path
        try:
            normalized_path = os.path.normpath(os.path.abspath(file_path))
            if not os.path.exists(normalized_path):
                return {"error": f"File not found: {file_path}"}
            if not os.path.isfile(normalized_path):
                return {"error": f"Path is not a file: {file_path}"}
            file_path = normalized_path
        except (OSError, ValueError) as e:
            return {"error": f"Invalid file path: {str(e)}"}

        # Get basic file info
        metadata = FileMetadataExtractor.get_basic_file_info(file_path)
        if "error" in metadata:
            return metadata

        extension = metadata["extension"]

        # Extract specific metadata based on file type
        specific_metadata = {}

        if extension in self.supported_image_formats:
            specific_metadata = FileMetadataExtractor.extract_image_metadata(file_path)
            metadata["file_type"] = "image"
        elif extension in self.supported_audio_formats:
            specific_metadata = FileMetadataExtractor.extract_audio_metadata(file_path)
            metadata["file_type"] = "audio"
        elif extension in self.supported_video_formats:
            specific_metadata = FileMetadataExtractor.extract_video_metadata(file_path)
            metadata["file_type"] = "video"
        elif extension == ".pdf":
            specific_metadata = FileMetadataExtractor.extract_pdf_metadata(file_path)
            metadata["file_type"] = "document"
        elif extension == ".docx":
            specific_metadata = FileMetadataExtractor.extract_docx_metadata(file_path)
            metadata["file_type"] = "document"
        else:
            metadata["file_type"] = "other"
            specific_metadata = {
                "note": "No specific metadata extraction available for this file type"
            }

        # Merge specific metadata
        if specific_metadata and "error" not in specific_metadata:
            metadata.update(specific_metadata)
        elif "error" in specific_metadata:
            metadata["extraction_error"] = specific_metadata["error"]

        return metadata

    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """
        Process all files in a directory and extract their metadata.

        Recursively walks through the directory structure and processes
        each file found, collecting metadata from all supported file types.

        Args:
            directory_path (str): Path to the directory to process

        Returns:
            List[Dict[str, Any]]: List of metadata dictionaries, one for each file
        """
        results = []

        for root, _, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                try:
                    metadata = self.extract_metadata(file_path)
                    results.append(metadata)
                except Exception as e:
                    results.append(
                        {
                            "filename": file,
                            "file_path": file_path,
                            "error": f"Failed to process: {str(e)}",
                        }
                    )

        return results

    @staticmethod
    def save_results(
        results: List[Dict[str, Any]], output_path: str, format_type: str = "json"
    ):
        """
        Save metadata extraction results to a file.

        Supports saving results in JSON or CSV format. JSON preserves the
        nested structure of metadata, while CSV flattens nested dictionaries.

        Args:
            results (List[Dict[str, Any]]): List of metadata dictionaries to save
            output_path (str): Path where the results file should be saved
            format_type (str): Output format, either 'json' or 'csv' (default: 'json')

        Note:
            CSV format flattens nested dictionaries using underscore notation
            (e.g., 'exif_Camera' for nested exif.Camera data).
        """
        try:
            if format_type.lower() == "json":
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(results, f, indent=2, ensure_ascii=False)

            elif format_type.lower() == "csv":
                if not results:
                    return

                # Get all unique keys for CSV headers
                all_keys = set()
                for result in results:
                    all_keys.update(result.keys())

                with open(output_path, "w", newline="", encoding="utf-8") as f:
                    writer = csv.DictWriter(f, fieldnames=sorted(all_keys))
                    writer.writeheader()

                    for result in results:
                        # Flatten nested dictionaries for CSV
                        flattened = {}
                        for key, value in result.items():
                            if isinstance(value, dict):
                                for sub_key, sub_value in value.items():
                                    flattened[f"{key}_{sub_key}"] = str(sub_value)
                            else:
                                flattened[key] = str(value) if value is not None else ""
                        writer.writerow(flattened)

            print(f"Results saved to: {output_path}")

        except Exception as e:
            print(f"Error saving results: {str(e)}")


def print_metadata(metadata: Dict[str, Any], indent: int = 0):
    """
    Pretty print metadata in a hierarchical format.

    Args:
        metadata (Dict[str, Any]): Dictionary containing metadata to print
        indent (int): Current indentation level for nested display
    """
    spacing = "  " * indent

    for key, value in metadata.items():
        if isinstance(value, dict):
            print(f"{spacing}{key}:")
            print_metadata(value, indent + 1)
        else:
            print(f"{spacing}{key}: {value}")


def main():
    """
    Main function to handle command line arguments and run the script.

    This function:
    - Parses command line arguments
    - Validates input paths
    - Processes files or directories
    - Handles output formatting and saving
    - Provides user feedback
    """
    parser = argparse.ArgumentParser(description="Extract metadata from files")
    parser.add_argument("path", help="File or directory path to analyze")
    parser.add_argument("-o", "--output", help="Output file path (optional)")
    parser.add_argument(
        "-f",
        "--format",
        choices=["json", "csv"],
        default="json",
        help="Output format (default: json)",
    )
    parser.add_argument(
        "-r", "--recursive", action="store_true", help="Process directories recursively"
    )

    args = parser.parse_args()

    extractor = FileMetadataExtractor()

    # Check if path exists
    if not os.path.exists(args.path):
        print(f"Error: Path '{args.path}' does not exist.")
        sys.exit(1)

    # Process single file or directory
    if os.path.isfile(args.path):
        print(f"Extracting metadata from: {args.path}")
        metadata = extractor.extract_metadata(args.path)

        if args.output:
            FileMetadataExtractor.save_results([metadata], args.output, args.format)
        else:
            print("\n--- Metadata ---")
            print_metadata(metadata)

    elif os.path.isdir(args.path):
        if args.recursive:
            print(f"Processing directory recursively: {args.path}")
            results = extractor.process_directory(args.path)
        else:
            print(f"Processing files in directory: {args.path}")
            results = []
            for file in os.listdir(args.path):
                file_path = os.path.join(args.path, file)
                if os.path.isfile(file_path):
                    metadata = extractor.extract_metadata(file_path)
                    results.append(metadata)

        if args.output:
            FileMetadataExtractor.save_results(results, args.output, args.format)
        else:
            print(f"\n--- Found {len(results)} files ---")
            for i, metadata in enumerate(results, 1):
                print(f"\n{i}. {metadata.get('filename', 'Unknown')}")
                print_metadata(metadata)
                if i >= 10:  # Limit console output
                    print(f"\n... and {len(results) - 10} more files")
                    break

    print("\nDone!")


if __name__ == "__main__":
    main()
